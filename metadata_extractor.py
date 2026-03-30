"""
Metadata Extractor Module
Extracts hidden metadata from various file types for forensic analysis.
Supports: JPG, PNG, PDF, DOCX

Security Features:
- No permanent file storage
- File type validation
- Size limits
- Sanitized outputs
"""

import os
import io
import tempfile
from typing import Dict, Any, Optional, Tuple, List
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum

# Image metadata libraries
try:
    import exifread
    EXIFREAD_AVAILABLE = True
except ImportError:
    EXIFREAD_AVAILABLE = False

try:
    from PIL import Image
    from PIL.ExifTags import TAGS, GPSTAGS
    PILLOW_AVAILABLE = True
except ImportError:
    PILLOW_AVAILABLE = False

# PDF metadata
try:
    from PyPDF2 import PdfReader
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

# DOCX metadata
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class FileType(Enum):
    """Supported file types for metadata extraction."""
    JPG = "image/jpeg"
    PNG = "image/png"
    PDF = "application/pdf"
    DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    UNKNOWN = "unknown"


@dataclass
class GPSCoordinates:
    """Represents GPS coordinates with human-readable format."""
    latitude: float
    longitude: float
    latitude_ref: str = "N"
    longitude_ref: str = "W"
    
    @property
    def decimal(self) -> Tuple[float, float]:
        """Return coordinates as decimal degrees."""
        lat = self.latitude if self.latitude_ref == "N" else -self.latitude
        lon = self.longitude if self.longitude_ref == "E" else -self.longitude
        return (lat, lon)
    
    @property
    def formatted(self) -> str:
        """Return human-readable coordinate string."""
        lat, lon = self.decimal
        lat_dir = "N" if lat >= 0 else "S"
        lon_dir = "E" if lon >= 0 else "W"
        return f"{abs(lat):.6f}° {lat_dir}, {abs(lon):.6f}° {lon_dir}"
    
    @property
    def maps_url(self) -> str:
        """Return Google Maps URL for the coordinates."""
        lat, lon = self.decimal
        return f"https://www.google.com/maps?q={lat},{lon}"
    
    @property
    def osm_url(self) -> str:
        """Return OpenStreetMap URL for the coordinates."""
        lat, lon = self.decimal
        return f"https://www.openstreetmap.org/?mlat={lat}&mlon={lon}&zoom=15"


@dataclass
class MetadataResult:
    """Stores extracted metadata with categorization."""
    success: bool
    file_type: str
    filename: str
    file_size: int
    metadata: Dict[str, Any] = field(default_factory=dict)
    gps_data: Optional[GPSCoordinates] = None
    error: Optional[str] = None
    extraction_time: str = field(default_factory=lambda: datetime.now().isoformat())
    
    # Categorized metadata
    camera_info: Dict[str, Any] = field(default_factory=dict)
    datetime_info: Dict[str, Any] = field(default_factory=dict)
    software_info: Dict[str, Any] = field(default_factory=dict)
    author_info: Dict[str, Any] = field(default_factory=dict)
    document_info: Dict[str, Any] = field(default_factory=dict)
    image_info: Dict[str, Any] = field(default_factory=dict)
    security_flags: List[str] = field(default_factory=list)


class MetadataExtractor:
    """
    Extracts metadata from various file types for forensic analysis.
    
    Security Features:
    - Maximum file size limit (default 50MB)
    - File type whitelist validation
    - Temporary file handling (no permanent storage)
    - Output sanitization
    """
    
    # Security constants
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
    ALLOWED_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.pdf', '.docx'}
    ALLOWED_MIMETYPES = {
        'image/jpeg': FileType.JPG,
        'image/png': FileType.PNG,
        'application/pdf': FileType.PDF,
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': FileType.DOCX
    }
    
    # File signatures (magic bytes) for validation
    FILE_SIGNATURES = {
        b'\xff\xd8\xff': FileType.JPG,      # JPEG
        b'\x89PNG\r\n\x1a\n': FileType.PNG,  # PNG
        b'%PDF': FileType.PDF,               # PDF
        b'PK\x03\x04': FileType.DOCX,        # DOCX (ZIP-based)
    }
    
    def __init__(self, max_file_size: int = None):
        """
        Initialize the metadata extractor.
        
        Args:
            max_file_size: Maximum allowed file size in bytes (default: 50MB)
        """
        self.max_file_size = max_file_size or self.MAX_FILE_SIZE
        self._check_dependencies()
    
    def _check_dependencies(self) -> Dict[str, bool]:
        """Check which optional dependencies are available."""
        return {
            'exifread': EXIFREAD_AVAILABLE,
            'pillow': PILLOW_AVAILABLE,
            'pypdf2': PYPDF2_AVAILABLE,
            'python-docx': DOCX_AVAILABLE
        }
    
    def get_available_features(self) -> Dict[str, bool]:
        """Return which file types can be processed based on installed dependencies."""
        deps = self._check_dependencies()
        return {
            'jpg': deps['exifread'] or deps['pillow'],
            'png': deps['pillow'],
            'pdf': deps['pypdf2'],
            'docx': deps['python-docx']
        }
    
    def validate_file(self, file_data: bytes, filename: str) -> Tuple[bool, str, FileType]:
        """
        Validate file for security before processing.
        
        Args:
            file_data: Raw file bytes
            filename: Original filename
            
        Returns:
            Tuple of (is_valid, error_message, file_type)
        """
        # Check file size
        if len(file_data) > self.max_file_size:
            return False, f"File too large. Maximum size: {self.max_file_size // (1024*1024)}MB", FileType.UNKNOWN
        
        # Check file extension
        ext = os.path.splitext(filename.lower())[1]
        if ext not in self.ALLOWED_EXTENSIONS:
            return False, f"Unsupported file type. Allowed: {', '.join(self.ALLOWED_EXTENSIONS)}", FileType.UNKNOWN
        
        # Validate file signature (magic bytes)
        file_type = self._detect_file_type(file_data)
        if file_type == FileType.UNKNOWN:
            return False, "Invalid file format or corrupted file", FileType.UNKNOWN
        
        # Cross-check extension with detected type
        ext_to_type = {
            '.jpg': FileType.JPG, '.jpeg': FileType.JPG,
            '.png': FileType.PNG,
            '.pdf': FileType.PDF,
            '.docx': FileType.DOCX
        }
        expected_type = ext_to_type.get(ext)
        if expected_type and expected_type != file_type:
            return False, f"File extension mismatch. Extension is {ext} but file content is {file_type.name}", FileType.UNKNOWN
        
        return True, "", file_type
    
    def _detect_file_type(self, file_data: bytes) -> FileType:
        """Detect file type based on magic bytes."""
        for signature, file_type in self.FILE_SIGNATURES.items():
            if file_data[:len(signature)] == signature:
                return file_type
        return FileType.UNKNOWN
    
    def _sanitize_value(self, value: Any) -> Any:
        """Sanitize metadata values to prevent injection attacks."""
        if isinstance(value, str):
            # Remove null bytes and control characters
            sanitized = ''.join(char for char in value if ord(char) >= 32 or char in '\n\r\t')
            # Limit string length
            return sanitized[:1000] if len(sanitized) > 1000 else sanitized
        elif isinstance(value, bytes):
            try:
                return self._sanitize_value(value.decode('utf-8', errors='ignore'))
            except:
                return "[Binary data]"
        elif isinstance(value, (int, float, bool)):
            return value
        elif isinstance(value, (list, tuple)):
            return [self._sanitize_value(v) for v in value[:100]]  # Limit list size
        elif isinstance(value, dict):
            return {str(k)[:100]: self._sanitize_value(v) for k, v in list(value.items())[:50]}
        else:
            return str(value)[:500]
    
    def extract(self, file_data: bytes, filename: str) -> MetadataResult:
        """
        Extract metadata from file.
        
        Args:
            file_data: Raw file bytes
            filename: Original filename
            
        Returns:
            MetadataResult with extracted metadata
        """
        # Validate file
        is_valid, error_msg, file_type = self.validate_file(file_data, filename)
        
        if not is_valid:
            return MetadataResult(
                success=False,
                file_type=file_type.name,
                filename=self._sanitize_value(filename),
                file_size=len(file_data),
                error=error_msg
            )
        
        # Extract based on file type
        try:
            if file_type == FileType.JPG:
                return self._extract_jpeg_metadata(file_data, filename)
            elif file_type == FileType.PNG:
                return self._extract_png_metadata(file_data, filename)
            elif file_type == FileType.PDF:
                return self._extract_pdf_metadata(file_data, filename)
            elif file_type == FileType.DOCX:
                return self._extract_docx_metadata(file_data, filename)
            else:
                return MetadataResult(
                    success=False,
                    file_type=file_type.name,
                    filename=self._sanitize_value(filename),
                    file_size=len(file_data),
                    error="Unsupported file type"
                )
        except Exception as e:
            return MetadataResult(
                success=False,
                file_type=file_type.name,
                filename=self._sanitize_value(filename),
                file_size=len(file_data),
                error=f"Extraction error: {str(e)}"
            )
    
    def _convert_to_degrees(self, value) -> float:
        """Convert GPS coordinates to decimal degrees."""
        if hasattr(value, 'values'):
            # exifread format
            d = float(value.values[0].num) / float(value.values[0].den)
            m = float(value.values[1].num) / float(value.values[1].den)
            s = float(value.values[2].num) / float(value.values[2].den)
        elif isinstance(value, (list, tuple)):
            # Pillow format
            d, m, s = value
            if hasattr(d, 'numerator'):
                d = float(d.numerator) / float(d.denominator)
                m = float(m.numerator) / float(m.denominator)
                s = float(s.numerator) / float(s.denominator)
            else:
                d, m, s = float(d), float(m), float(s)
        else:
            return float(value)
        
        return d + (m / 60.0) + (s / 3600.0)
    
    def _extract_gps_from_exif(self, tags: Dict) -> Optional[GPSCoordinates]:
        """Extract GPS coordinates from EXIF tags."""
        try:
            # Try exifread format
            if 'GPS GPSLatitude' in tags and 'GPS GPSLongitude' in tags:
                lat = self._convert_to_degrees(tags['GPS GPSLatitude'])
                lon = self._convert_to_degrees(tags['GPS GPSLongitude'])
                lat_ref = str(tags.get('GPS GPSLatitudeRef', 'N'))
                lon_ref = str(tags.get('GPS GPSLongitudeRef', 'W'))
                return GPSCoordinates(lat, lon, lat_ref, lon_ref)
            
            # Try Pillow GPSInfo format
            if 'GPSInfo' in tags:
                gps_info = tags['GPSInfo']
                if isinstance(gps_info, dict):
                    lat = gps_info.get(2) or gps_info.get('GPSLatitude')
                    lon = gps_info.get(4) or gps_info.get('GPSLongitude')
                    lat_ref = gps_info.get(1, 'N') or gps_info.get('GPSLatitudeRef', 'N')
                    lon_ref = gps_info.get(3, 'W') or gps_info.get('GPSLongitudeRef', 'W')
                    
                    if lat and lon:
                        return GPSCoordinates(
                            self._convert_to_degrees(lat),
                            self._convert_to_degrees(lon),
                            str(lat_ref),
                            str(lon_ref)
                        )
        except Exception:
            pass
        return None
    
    def _extract_jpeg_metadata(self, file_data: bytes, filename: str) -> MetadataResult:
        """Extract metadata from JPEG files using exifread and Pillow."""
        result = MetadataResult(
            success=True,
            file_type="JPEG",
            filename=self._sanitize_value(filename),
            file_size=len(file_data)
        )
        
        all_metadata = {}
        
        # Use exifread for detailed EXIF data
        if EXIFREAD_AVAILABLE:
            try:
                file_stream = io.BytesIO(file_data)
                tags = exifread.process_file(file_stream, details=True)
                
                for tag, value in tags.items():
                    if tag.startswith('Thumbnail'):
                        continue  # Skip thumbnail data
                    tag_name = str(tag)
                    all_metadata[tag_name] = self._sanitize_value(str(value))
                    
                    # Categorize metadata
                    if 'Camera' in tag_name or 'Make' in tag_name or 'Model' in tag_name:
                        result.camera_info[tag_name] = self._sanitize_value(str(value))
                    elif 'Date' in tag_name or 'Time' in tag_name:
                        result.datetime_info[tag_name] = self._sanitize_value(str(value))
                    elif 'Software' in tag_name or 'Processing' in tag_name:
                        result.software_info[tag_name] = self._sanitize_value(str(value))
                    elif 'Artist' in tag_name or 'Author' in tag_name or 'Copyright' in tag_name:
                        result.author_info[tag_name] = self._sanitize_value(str(value))
                
                # Extract GPS
                gps = self._extract_gps_from_exif(tags)
                if gps:
                    result.gps_data = gps
                    result.security_flags.append("GPS_LOCATION_FOUND")
                    
            except Exception as e:
                all_metadata['_exifread_error'] = str(e)
        
        # Use Pillow for additional info
        if PILLOW_AVAILABLE:
            try:
                file_stream = io.BytesIO(file_data)
                with Image.open(file_stream) as img:
                    result.image_info['format'] = img.format
                    result.image_info['mode'] = img.mode
                    result.image_info['width'] = img.width
                    result.image_info['height'] = img.height
                    result.image_info['resolution'] = f"{img.width}x{img.height}"
                    
                    # Get EXIF from Pillow if exifread failed
                    if not all_metadata:
                        exif = img._getexif()
                        if exif:
                            for tag_id, value in exif.items():
                                tag = TAGS.get(tag_id, tag_id)
                                if tag == 'GPSInfo':
                                    gps_data = {}
                                    for gps_tag_id, gps_value in value.items():
                                        gps_tag = GPSTAGS.get(gps_tag_id, gps_tag_id)
                                        gps_data[gps_tag] = self._sanitize_value(gps_value)
                                    all_metadata['GPSInfo'] = gps_data
                                    
                                    # Try to extract GPS if not already found
                                    if not result.gps_data:
                                        gps = self._extract_gps_from_exif({'GPSInfo': value})
                                        if gps:
                                            result.gps_data = gps
                                            result.security_flags.append("GPS_LOCATION_FOUND")
                                else:
                                    all_metadata[str(tag)] = self._sanitize_value(value)
                                    
            except Exception as e:
                all_metadata['_pillow_error'] = str(e)
        
        result.metadata = all_metadata
        
        # Add security flags
        if result.camera_info:
            result.security_flags.append("CAMERA_INFO_FOUND")
        if result.author_info:
            result.security_flags.append("AUTHOR_INFO_FOUND")
        if result.software_info:
            result.security_flags.append("SOFTWARE_INFO_FOUND")
            
        return result
    
    def _extract_png_metadata(self, file_data: bytes, filename: str) -> MetadataResult:
        """Extract metadata from PNG files using Pillow."""
        result = MetadataResult(
            success=True,
            file_type="PNG",
            filename=self._sanitize_value(filename),
            file_size=len(file_data)
        )
        
        if not PILLOW_AVAILABLE:
            result.success = False
            result.error = "Pillow library not installed. Install with: pip install Pillow"
            return result
        
        all_metadata = {}
        
        try:
            file_stream = io.BytesIO(file_data)
            with Image.open(file_stream) as img:
                result.image_info['format'] = img.format
                result.image_info['mode'] = img.mode
                result.image_info['width'] = img.width
                result.image_info['height'] = img.height
                result.image_info['resolution'] = f"{img.width}x{img.height}"
                
                # PNG metadata is stored in info dict
                if img.info:
                    for key, value in img.info.items():
                        if key not in ('exif', 'icc_profile'):  # Skip binary data
                            all_metadata[key] = self._sanitize_value(value)
                            
                            if 'Software' in str(key) or 'Creator' in str(key):
                                result.software_info[key] = self._sanitize_value(value)
                            elif 'Author' in str(key) or 'Artist' in str(key):
                                result.author_info[key] = self._sanitize_value(value)
                            elif 'Date' in str(key) or 'Time' in str(key):
                                result.datetime_info[key] = self._sanitize_value(value)
                
                # Check for EXIF in PNG
                if hasattr(img, '_getexif') and img._getexif():
                    exif = img._getexif()
                    for tag_id, value in exif.items():
                        tag = TAGS.get(tag_id, tag_id)
                        all_metadata[f"EXIF_{tag}"] = self._sanitize_value(value)
                        
        except Exception as e:
            result.success = False
            result.error = f"PNG extraction error: {str(e)}"
            
        result.metadata = all_metadata
        
        if result.software_info:
            result.security_flags.append("SOFTWARE_INFO_FOUND")
        if result.author_info:
            result.security_flags.append("AUTHOR_INFO_FOUND")
            
        return result
    
    def _extract_pdf_metadata(self, file_data: bytes, filename: str) -> MetadataResult:
        """Extract metadata from PDF files using PyPDF2."""
        result = MetadataResult(
            success=True,
            file_type="PDF",
            filename=self._sanitize_value(filename),
            file_size=len(file_data)
        )
        
        if not PYPDF2_AVAILABLE:
            result.success = False
            result.error = "PyPDF2 library not installed. Install with: pip install PyPDF2"
            return result
        
        all_metadata = {}
        
        try:
            file_stream = io.BytesIO(file_data)
            reader = PdfReader(file_stream)
            
            # Document info
            result.document_info['pages'] = len(reader.pages)
            result.document_info['encrypted'] = reader.is_encrypted
            
            all_metadata['page_count'] = len(reader.pages)
            all_metadata['is_encrypted'] = reader.is_encrypted
            
            if reader.is_encrypted:
                result.security_flags.append("DOCUMENT_ENCRYPTED")
            
            # PDF metadata
            if reader.metadata:
                metadata = reader.metadata
                
                # Map PDF metadata fields
                field_mapping = {
                    '/Title': ('title', 'document_info'),
                    '/Author': ('author', 'author_info'),
                    '/Subject': ('subject', 'document_info'),
                    '/Creator': ('creator', 'software_info'),
                    '/Producer': ('producer', 'software_info'),
                    '/CreationDate': ('creation_date', 'datetime_info'),
                    '/ModDate': ('modification_date', 'datetime_info'),
                    '/Keywords': ('keywords', 'document_info'),
                }
                
                for pdf_key, (clean_key, category) in field_mapping.items():
                    value = getattr(metadata, pdf_key.replace('/', '').lower(), None)
                    if value:
                        sanitized = self._sanitize_value(value)
                        all_metadata[clean_key] = sanitized
                        
                        if category == 'author_info':
                            result.author_info[clean_key] = sanitized
                        elif category == 'software_info':
                            result.software_info[clean_key] = sanitized
                        elif category == 'datetime_info':
                            result.datetime_info[clean_key] = sanitized
                        else:
                            result.document_info[clean_key] = sanitized
                
                # Check for custom metadata
                # Skip deprecated/removed properties in PyPDF2 3.0+
                skip_properties = {'xmpMetadata', 'xmp_metadata'}
                for key in dir(metadata):
                    if not key.startswith('_') and key not in skip_properties:
                        try:
                            value = getattr(metadata, key)
                            if value and callable(value) is False:
                                all_metadata[key] = self._sanitize_value(value)
                        except (AttributeError, DeprecationWarning):
                            # Skip deprecated or removed attributes
                            pass
            
            # Extract XMP metadata (PyPDF2 3.0+ uses xmp_metadata)
            try:
                xmp = reader.xmp_metadata
                if xmp:
                    xmp_data = {}
                    # Common XMP properties
                    xmp_props = [
                        ('dc_title', 'title'),
                        ('dc_creator', 'creator'),
                        ('dc_description', 'description'),
                        ('dc_subject', 'subject'),
                        ('xmp_create_date', 'create_date'),
                        ('xmp_modify_date', 'modify_date'),
                        ('xmp_creator_tool', 'creator_tool'),
                        ('pdf_producer', 'producer'),
                    ]
                    for attr, key in xmp_props:
                        try:
                            value = getattr(xmp, attr, None)
                            if value:
                                xmp_data[key] = self._sanitize_value(str(value))
                        except Exception:
                            pass
                    if xmp_data:
                        all_metadata['xmp_metadata'] = xmp_data
            except Exception:
                # XMP metadata not available or error reading it
                pass
                            
        except Exception as e:
            result.success = False
            result.error = f"PDF extraction error: {str(e)}"
            
        result.metadata = all_metadata
        
        if result.author_info:
            result.security_flags.append("AUTHOR_INFO_FOUND")
        if result.software_info:
            result.security_flags.append("SOFTWARE_INFO_FOUND")
            
        return result
    
    def _extract_docx_metadata(self, file_data: bytes, filename: str) -> MetadataResult:
        """Extract metadata from DOCX files using python-docx."""
        result = MetadataResult(
            success=True,
            file_type="DOCX",
            filename=self._sanitize_value(filename),
            file_size=len(file_data)
        )
        
        if not DOCX_AVAILABLE:
            result.success = False
            result.error = "python-docx library not installed. Install with: pip install python-docx"
            return result
        
        all_metadata = {}
        
        try:
            file_stream = io.BytesIO(file_data)
            doc = Document(file_stream)
            
            # Core properties
            core_props = doc.core_properties
            
            # Author info
            if core_props.author:
                result.author_info['author'] = self._sanitize_value(core_props.author)
                all_metadata['author'] = self._sanitize_value(core_props.author)
            if core_props.last_modified_by:
                result.author_info['last_modified_by'] = self._sanitize_value(core_props.last_modified_by)
                all_metadata['last_modified_by'] = self._sanitize_value(core_props.last_modified_by)
            
            # Document info
            if core_props.title:
                result.document_info['title'] = self._sanitize_value(core_props.title)
                all_metadata['title'] = self._sanitize_value(core_props.title)
            if core_props.subject:
                result.document_info['subject'] = self._sanitize_value(core_props.subject)
                all_metadata['subject'] = self._sanitize_value(core_props.subject)
            if core_props.keywords:
                result.document_info['keywords'] = self._sanitize_value(core_props.keywords)
                all_metadata['keywords'] = self._sanitize_value(core_props.keywords)
            if core_props.comments:
                result.document_info['comments'] = self._sanitize_value(core_props.comments)
                all_metadata['comments'] = self._sanitize_value(core_props.comments)
            if core_props.category:
                result.document_info['category'] = self._sanitize_value(core_props.category)
                all_metadata['category'] = self._sanitize_value(core_props.category)
            if core_props.content_status:
                result.document_info['content_status'] = self._sanitize_value(core_props.content_status)
                all_metadata['content_status'] = self._sanitize_value(core_props.content_status)
            
            # Datetime info
            if core_props.created:
                result.datetime_info['created'] = core_props.created.isoformat() if core_props.created else None
                all_metadata['created'] = core_props.created.isoformat() if core_props.created else None
            if core_props.modified:
                result.datetime_info['modified'] = core_props.modified.isoformat() if core_props.modified else None
                all_metadata['modified'] = core_props.modified.isoformat() if core_props.modified else None
            if core_props.last_printed:
                result.datetime_info['last_printed'] = core_props.last_printed.isoformat() if core_props.last_printed else None
                all_metadata['last_printed'] = core_props.last_printed.isoformat() if core_props.last_printed else None
            
            # Revision info
            if core_props.revision:
                result.document_info['revision'] = core_props.revision
                all_metadata['revision'] = core_props.revision
            if core_props.version:
                result.document_info['version'] = self._sanitize_value(core_props.version)
                all_metadata['version'] = self._sanitize_value(core_props.version)
            
            # Document statistics
            result.document_info['paragraph_count'] = len(doc.paragraphs)
            result.document_info['table_count'] = len(doc.tables)
            result.document_info['section_count'] = len(doc.sections)
            
            all_metadata['paragraph_count'] = len(doc.paragraphs)
            all_metadata['table_count'] = len(doc.tables)
            all_metadata['section_count'] = len(doc.sections)
            
        except Exception as e:
            result.success = False
            result.error = f"DOCX extraction error: {str(e)}"
            
        result.metadata = all_metadata
        
        if result.author_info:
            result.security_flags.append("AUTHOR_INFO_FOUND")
        if result.datetime_info:
            result.security_flags.append("DATETIME_INFO_FOUND")
            
        return result
    
    def to_dict(self, result: MetadataResult) -> Dict[str, Any]:
        """Convert MetadataResult to dictionary for JSON serialization."""
        data = {
            'success': result.success,
            'file_type': result.file_type,
            'filename': result.filename,
            'file_size': result.file_size,
            'file_size_formatted': self._format_file_size(result.file_size),
            'metadata': result.metadata,
            'extraction_time': result.extraction_time,
            'camera_info': result.camera_info,
            'datetime_info': result.datetime_info,
            'software_info': result.software_info,
            'author_info': result.author_info,
            'document_info': result.document_info,
            'image_info': result.image_info,
            'security_flags': result.security_flags,
            'error': result.error
        }
        
        if result.gps_data:
            lat, lon = result.gps_data.decimal
            data['gps_data'] = {
                'latitude': lat,
                'longitude': lon,
                'formatted': result.gps_data.formatted,
                'maps_url': result.gps_data.maps_url,
                'osm_url': result.gps_data.osm_url
            }
        else:
            data['gps_data'] = None
            
        return data
    
    def _format_file_size(self, size_bytes: int) -> str:
        """Format file size to human-readable string."""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size_bytes < 1024:
                return f"{size_bytes:.2f} {unit}"
            size_bytes /= 1024
        return f"{size_bytes:.2f} TB"


# Convenience function for quick extraction
def extract_metadata(file_data: bytes, filename: str) -> Dict[str, Any]:
    """
    Quick metadata extraction function.
    
    Args:
        file_data: Raw file bytes
        filename: Original filename
        
    Returns:
        Dictionary with extracted metadata
    """
    extractor = MetadataExtractor()
    result = extractor.extract(file_data, filename)
    return extractor.to_dict(result)


if __name__ == "__main__":
    # Test the module
    extractor = MetadataExtractor()
    
    print("=" * 60)
    print("Metadata Extractor - Dependency Check")
    print("=" * 60)
    
    deps = extractor._check_dependencies()
    for dep, available in deps.items():
        status = "✓ Installed" if available else "✗ Not installed"
        print(f"  {dep}: {status}")
    
    print("\nSupported file types:")
    features = extractor.get_available_features()
    for file_type, supported in features.items():
        status = "✓ Supported" if supported else "✗ Missing dependencies"
        print(f"  {file_type.upper()}: {status}")
    
    print("\n" + "=" * 60)
    print("Ready for integration with Dark Web Leak Monitor")
    print("=" * 60)
